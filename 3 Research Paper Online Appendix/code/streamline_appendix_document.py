"""Replace embedded quantitative appendices with concise GitHub pointers.

The script preserves the manuscript and all non-quantitative appendix content.
It keeps the Appendix A.1 and C-G titles/subtitles, removes the tables and
figures stored in the public online appendix, and inserts one short descriptive
paragraph with a direct GitHub hyperlink for each quantitative subsection.

No index, regression, interaction, latent-class model, posterior assignment,
segment definition, or substantive result is estimated or modified here.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt, RGBColor


REPOSITORY_ROOT = (
    "https://github.com/Fundacion-Capital/"
    "CFI-DPI-Migrant-Women-Colombia/tree/main/"
    "3%20Research%20Paper%20Online%20Appendix"
)

POINTERS = [
    {
        "start": "A.1 Full quantitative survey questionnaire administered through SurveyCTO/CATI",
        "end": "A.2 In-depth interview guide for Venezuelan migrant women",
        "description": (
            "The complete CATI/SurveyCTO questionnaire covers consent, eligibility, "
            "demographics, migration, digital access, financial use, remittances, "
            "fraud and recourse, autonomy, and enabling support."
        ),
        "url": f"{REPOSITORY_ROOT}/A%20Research%20Instrument",
    },
    {
        "start": "C.1 Index architecture",
        "end": "C.2 Item-to-index mapping",
        "description": (
            "This subsection documents each quantitative index's construct, components, "
            "aggregation rule, scale, direction, and descriptive cut points."
        ),
        "url": f"{REPOSITORY_ROOT}/C%20Variable%20Construction/tables",
    },
    {
        "start": "C.2 Item-to-index mapping",
        "end": "C.3 Index distributions and associations",
        "description": (
            "This subsection maps each survey item and constructed component to its "
            "implemented scoring rule and role in the corresponding index."
        ),
        "url": f"{REPOSITORY_ROOT}/C%20Variable%20Construction/tables",
    },
    {
        "start": "C.3 Index distributions and associations",
        "end": "Appendix D. Quantitative descriptive evidence",
        "description": (
            "This subsection reports aggregate index summary statistics, pairwise "
            "correlations, distribution plots, and the index-correlation heatmap."
        ),
        "url": f"{REPOSITORY_ROOT}/C%20Variable%20Construction",
    },
    {
        "start": "D.1 Extended descriptive tables",
        "end": "D.2 Extended descriptive figures",
        "description": (
            "This subsection provides extended categorical distributions and "
            "multiple-response prevalence tables for the quantitative sample."
        ),
        "url": f"{REPOSITORY_ROOT}/D%20Descriptive%20Evidence/tables",
    },
    {
        "start": "D.2 Extended descriptive figures",
        "end": "Appendix E. Multivariate regression analysis",
        "description": (
            "This subsection provides the full descriptive figure gallery covering "
            "respondent characteristics, access, digital competence, financial use, "
            "remittances, fraud, autonomy, trust, barriers, and enabling support."
        ),
        "url": f"{REPOSITORY_ROOT}/D%20Descriptive%20Evidence/figures",
    },
    {
        "start": "E.1 Complete regression tables",
        "end": "E.2 Coefficient, marginal-effect, and interaction figures",
        "description": (
            "This subsection provides the complete multivariate regression tables, "
            "formal-remittance model, interaction specifications, and supplementary "
            "item-level models."
        ),
        "url": f"{REPOSITORY_ROOT}/E%20Multivariate%20Regressions/tables",
    },
    {
        "start": "E.2 Coefficient, marginal-effect, and interaction figures",
        "end": "Appendix F. Latent class analysis diagnostics and model selection",
        "description": (
            "This subsection provides all coefficient plots, average-marginal-effect "
            "figures, and interaction plots associated with the reported models."
        ),
        "url": f"{REPOSITORY_ROOT}/E%20Multivariate%20Regressions/figures",
    },
    {
        "start": "F.1 Diagnostic and model-selection tables",
        "end": "F.2 Indicator-distribution and classification-diagnostic figures",
        "description": (
            "This subsection provides sample checks, indicator inventories, missingness "
            "and recoding diagnostics, feature sets, model-selection evidence, fit "
            "statistics, response patterns, and classification-certainty tables."
        ),
        "url": f"{REPOSITORY_ROOT}/F%20LCA%20Diagnostics/tables",
    },
    {
        "start": "F.2 Indicator-distribution and classification-diagnostic figures",
        "end": "Appendix G. Final LCA segment profiles and post-LCA validation",
        "description": (
            "This subsection provides the complete LCA diagnostic gallery, including "
            "index distributions, category distributions, association plots, class "
            "profiles, segment sizes, and posterior-classification diagnostics."
        ),
        "url": f"{REPOSITORY_ROOT}/F%20LCA%20Diagnostics/figures",
    },
    {
        "start": "G.1 Final segment tables",
        "end": "G.2 Segment-profile and post-LCA figures",
        "description": (
            "This subsection provides final class labels, conditional probabilities, "
            "segment size and certainty, centroids, demographic and behavioral profiles, "
            "multinomial models, outcome validation, and the policy typology."
        ),
        "url": f"{REPOSITORY_ROOT}/G%20Segment%20Profiles/tables",
    },
    {
        "start": "G.2 Segment-profile and post-LCA figures",
        "end": "Appendix H. Qualitative coding, segmentation, and evidence matrices",
        "description": (
            "This subsection provides the final class-profile probabilities, standardized "
            "centroids, segment sizes, posterior probabilities and certainty, and the "
            "class-defining and auxiliary outcome figures."
        ),
        "url": f"{REPOSITORY_ROOT}/G%20Segment%20Profiles/figures",
    },
]


def element_text(element) -> str:
    # ``itertext()`` repeats text for python-docx's nested run proxies.  Read
    # only Word text nodes so markers are compared to the visible paragraph.
    return "".join(node.text or "" for node in element.iter(qn("w:t"))).strip()


def body_paragraphs(document: Document) -> dict[str, object]:
    found: dict[str, object] = {}
    for child in document._element.body:
        if child.tag == qn("w:p"):
            text = element_text(child)
            if text:
                found[text] = child
    return found


def set_outline_body_text(paragraph_element) -> None:
    properties = paragraph_element.get_or_add_pPr()
    for existing in properties.findall(qn("w:outlineLvl")):
        properties.remove(existing)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "9")
    properties.append(outline)


def add_hyperlink(paragraph, url: str, text: str) -> None:
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)

    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    properties.append(size)
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build_pointer_paragraph(document: Document, description: str, url: str):
    paragraph = document.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(description + " ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(31, 41, 51)
    label = paragraph.add_run("Online materials: ")
    label.bold = True
    label.font.name = "Times New Roman"
    label.font.size = Pt(11)
    add_hyperlink(paragraph, url, url)
    paragraph.add_run(".")
    set_outline_body_text(paragraph._p)
    return paragraph._p


def replace_quantitative_blocks(document: Document) -> list[dict[str, object]]:
    body = document._element.body
    paragraph_lookup = body_paragraphs(document)
    missing = [item["start"] for item in POINTERS if item["start"] not in paragraph_lookup]
    missing += [item["end"] for item in POINTERS if item["end"] not in paragraph_lookup]
    if missing:
        raise RuntimeError("Required appendix markers were not found: " + "; ".join(sorted(set(missing))))

    indexed = []
    children = list(body)
    for item in POINTERS:
        start = paragraph_lookup[item["start"]]
        end = paragraph_lookup[item["end"]]
        start_index = children.index(start)
        end_index = children.index(end)
        if end_index <= start_index:
            raise RuntimeError(f"Invalid marker order: {item['start']} -> {item['end']}")
        indexed.append((start_index, item, start, end))

    report = []
    for _, item, start, end in sorted(indexed, key=lambda value: value[0], reverse=True):
        current = list(body)
        start_index = current.index(start)
        end_index = current.index(end)
        removed = current[start_index + 1 : end_index]
        removed_tables = sum(element.tag == qn("w:tbl") for element in removed)
        removed_drawings = sum(len(element.xpath(".//w:drawing")) for element in removed)
        for element in removed:
            body.remove(element)

        pointer = build_pointer_paragraph(document, item["description"], item["url"])
        end.addprevious(pointer)
        report.append(
            {
                "subsection": item["start"],
                "removed_body_elements": len(removed),
                "removed_tables": removed_tables,
                "removed_drawings": removed_drawings,
                "url": item["url"],
            }
        )
    return list(reversed(report))


def enforce_appendix_navigation(document: Document) -> None:
    appendix_started = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text == "Appendix":
            paragraph.style = document.styles["Heading 1"]
            appendix_started = True
            continue
        if not appendix_started:
            continue
        if paragraph.style.name.startswith("Heading"):
            paragraph.style = document.styles["Normal"]
        set_outline_body_text(paragraph._p)


def prune_unused_document_images(document: Document) -> int:
    used = set()
    for element in document._element.xpath(".//*[@r:embed]"):
        relation_id = element.get(qn("r:embed"))
        if relation_id:
            used.add(relation_id)
    removed = 0
    for relation_id, relationship in list(document.part.rels.items()):
        if relationship.reltype == RT.IMAGE and relation_id not in used:
            document.part.drop_rel(relation_id)
            removed += 1
    return removed


def enable_field_updates(document: Document) -> None:
    settings = document.settings.element
    for existing in settings.findall(qn("w:updateFields")):
        settings.remove(existing)
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def count_appendix_internal_headings(document: Document) -> int:
    started = False
    count = 0
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "Appendix" and paragraph.style.name == "Heading 1":
            started = True
            continue
        if started and paragraph.style.name.startswith("Heading"):
            count += 1
    return count


def count_online_appendix_hyperlinks(document: Document) -> int:
    """Count hyperlink elements, not deduplicated relationship targets."""
    count = 0
    for hyperlink in document._element.xpath(".//w:hyperlink[@r:id]"):
        relation_id = hyperlink.get(qn("r:id"))
        relationship = document.part.rels.get(relation_id)
        if (
            relationship is not None
            and relationship.reltype == RT.HYPERLINK
            and "tree/main/3%20Research%20Paper%20Online%20Appendix"
            in relationship.target_ref
        ):
            count += 1
    return count


def streamline(source: Path, output: Path, report_path: Path | None) -> dict[str, object]:
    document = Document(source)
    before = {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "bytes": source.stat().st_size,
    }
    replacements = replace_quantitative_blocks(document)
    enforce_appendix_navigation(document)
    pruned_images = prune_unused_document_images(document)
    enable_field_updates(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)

    verified = Document(output)
    after = {
        "paragraphs": len(verified.paragraphs),
        "tables": len(verified.tables),
        "inline_shapes": len(verified.inline_shapes),
        "bytes": output.stat().st_size,
        "appendix_internal_heading_styles": count_appendix_internal_headings(verified),
        "github_hyperlinks": count_online_appendix_hyperlinks(verified),
    }
    report = {
        "source": str(source),
        "output": str(output),
        "before": before,
        "after": after,
        "pruned_unused_image_relationships": pruned_images,
        "replacements": replacements,
    }
    if after["appendix_internal_heading_styles"] != 0:
        raise RuntimeError("Appendix-internal Heading styles remain after streamlining")
    if after["github_hyperlinks"] != len(POINTERS):
        raise RuntimeError("Not all quantitative appendix hyperlinks were created")
    if report_path is not None:
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(json.dumps(report, indent=2), encoding="utf-8")
    return report


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--report", type=Path)
    return parser.parse_args()


if __name__ == "__main__":
    arguments = parse_args()
    result = streamline(arguments.source, arguments.output, arguments.report)
    print(json.dumps(result, indent=2))
