"""Polish the complete appendix section of the integrated working paper.

The script is intentionally editorial. It does not estimate or change any
index, regression, interaction, latent-class model, posterior assignment, or
segment definition. It converts appendix-internal headings to non-TOC styles,
restores a visible title hierarchy, embeds the two missing workbook tables as
native Word tables, standardizes appendix typography, and marks fields for a
Word refresh.
"""

from __future__ import annotations

import argparse
import json
import re
from copy import deepcopy
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from openpyxl.utils.cell import range_boundaries


REPO_ROOT = Path(__file__).resolve().parents[2]
APPENDIX_ROOT = REPO_ROOT / "3 Research Paper Online Appendix"
ARCHITECTURE_WORKBOOK = (
    APPENDIX_ROOT
    / "C Variable Construction"
    / "tables"
    / "Table_C0_Current_Index_Architecture.xlsx"
)

APPENDIX_TITLE_STYLE = "Appendix Title"
APPENDIX_SECTION_STYLE = "Appendix Section"
APPENDIX_SUBHEADING_STYLE = "Appendix Subheading"
APPENDIX_MINOR_STYLE = "Appendix Minor Heading"
APPENDIX_NOTE_STYLE = "Appendix Note"

TEAL = "0F4C5C"
PALE_GRAY = "F4F6F7"
TEXT_HEX = "1F2933"
WHITE_HEX = "FFFFFF"
TEXT = RGBColor(31, 41, 51)

F_TABLE_TITLES = {
    1: "Sample integrity",
    2: "Indicator inventory",
    3: "Missingness and unique values",
    4: "Category distributions",
    5: "Pairwise associations",
    6: "LCA recoding",
    7: "Feature sets",
    8: "Model selection",
    9: "Final LCA fit",
    10: "Posterior certainty",
    11: "Minimal response patterns",
    12: "Minimal fit summary",
    13: "Minimal class profiles",
    14: "Hybrid distribution diagnostics",
    15: "Hybrid category distributions",
    16: "Hybrid response patterns",
    17: "Hybrid fit summary",
    18: "Hybrid class profiles",
}

G_TABLE_TITLES = {
    1: "Class mapping and labels",
    2: "Conditional response probabilities",
    3: "Segment size and certainty",
    4: "Continuous-index centroids",
    5: "Signature behaviors",
    6: "Demographic profile",
    7: "Financial-behavior profile",
    8: "Risk, autonomy, and support",
    9: "Core outcomes",
    10: "Multinomial predictors",
    11: "Average marginal effects",
    12: "Multinomial joint tests",
    13: "Segment-outcome regressions",
    14: "Segment-outcome joint tests",
    15: "Full class profiles",
    16: "Segment centroids",
    17: "Policy typology",
}

SHEET_TITLES = {
    "variable_inventory": "Variable inventory",
    "methodological_rules": "Methodological rules",
    "candidate_inputs_summary": "Candidate-input summary",
    "continuous_correlations_long": "Long-format continuous correlations",
    "continuous_correlation_matrix": "Continuous correlation matrix",
    "categorical_cramersv": "Cramer's V for categorical variables",
    "flagged_redundancy_pairs": "Flagged redundancy pairs",
    "recoding_metadata": "Recoding metadata",
    "recoded_distributions": "Recoded distributions",
    "feature_sets_overview": "Feature-set overview",
    "feature_set_variables": "Feature-set variables",
    "feature_set_diagnostics": "Feature-set diagnostics",
    "all_models": "All candidate models",
    "converged_models": "Converged models",
    "report_candidates": "Models retained for reporting",
    "selected_model": "Selected model",
    "by_class": "By latent class",
    "overall": "Overall summary",
    "by_segment": "By named segment",
    "raw_to_segment_mapping": "Raw-class to segment mapping",
    "profile_audit": "Profile audit",
    "long": "Long format",
    "posterior_weighted_wide": "Posterior-weighted wide format",
    "modal_assignment_wide": "Modal-assignment wide format",
    "profile_summary": "Profile summary",
    "wide": "Wide format",
    "binary_long": "Binary indicators, long format",
    "binary_wide": "Binary indicators, wide format",
    "categorical_long": "Categorical indicators, long format",
    "categorical_wide": "Categorical indicators, wide format",
    "compact_report_profile": "Compact reporting profile",
    "segment_sizes": "Segment sizes",
    "cat_long": "Categorical variables, long format",
    "cat_report": "Categorical reporting table",
    "cat_tests": "Categorical omnibus tests",
    "cont_long": "Continuous variables, long format",
    "cont_report": "Continuous reporting table",
    "cont_tests": "Continuous omnibus tests",
    "pairwise_holm": "Holm-adjusted pairwise comparisons",
    "outcome_definitions": "Outcome definitions",
    "outcome_long": "Outcomes, long format",
    "outcome_report": "Outcome reporting table",
    "outcome_tests": "Outcome omnibus tests",
    "external_validation": "External validation",
    "class_defining": "Class-defining outcomes",
    "figure_C10_data": "Data underlying Figure C.10",
    "segment_summary": "Segment summary",
    "full_categories_long": "Full categories, long format",
    "full_categories_wide": "Full categories, wide format",
    "centroids_long": "Centroids, long format",
    "centroids_wide": "Centroids, wide format",
    "report_centroids": "Reporting centroids",
    "policy_typology": "Policy typology",
    "figure_manifest": "Figure manifest",
}

F_FIGURE_TITLES = {
    61: "Outcome means by segment",
    62: "Outcome means by segment (presentation format)",
    63: "Class-defining outcomes",
    64: "External outcomes and mechanisms",
    65: "Index distributions",
    66: "Index histograms",
    67: "Correlation heatmap",
    68: "Final class-profile probabilities",
    69: "Standardized final-class centroids",
    70: "Final class sizes",
    71: "Posterior certainty",
    72: "Distribution of posterior probabilities",
    73: "Posterior certainty by segment",
}

G_FIGURE_TITLES = {
    1: "Class-profile probabilities",
    2: "Standardized segment centroids",
    3: "Segment sizes",
    4: "Distribution of posterior probabilities",
    5: "Posterior certainty by segment",
    6: "Class-defining outcomes",
    7: "Auxiliary outcomes and mechanisms",
}

HEADING_REPLACEMENTS = {
    "H.1 Qualitative Coding Framework": "H.1 Qualitative coding framework",
    "H.2 Full Qualitative Code Matrix": "H.2 Full qualitative code matrix",
    "H.3 Qualitative User Position Matrix": "H.3 Qualitative user-position matrix",
    "H.4 Evidence Matrix: Qualitative Themes and Analytical Mechanisms": "H.4 Evidence matrix: qualitative themes and analytical mechanisms",
    "H.5 Qualitative Integration Matrix": "H.5 Qualitative integration matrix",
    "Table H1. Qualitative Codebook": "Table H.1. Qualitative codebook",
    "Table H2. Main Themes and Associated Codes": "Table H.2. Main themes and associated codes",
    "Table H3. Qualitative User Positions": "Table H.3. Qualitative user positions",
    "Table H4. Qualitative Evidence Matrix": "Table H.4. Qualitative evidence matrix",
    "Table H5. From Codes to Analytical Findings": "Table H.5. From codes to analytical findings",
    "I.1 Purpose of the Benchmark": "I.1 Purpose of the benchmark",
    "I.2 DPI Map Extraction Methodology": "I.2 DPI map extraction methodology",
    "Countries Included": "Countries included",
    "DPI Pillars Considered": "DPI pillars considered",
    "Variables Selected": "Variables selected",
    "Treatment of Missing Information": "Treatment of missing information",
    "I.3 Benchmark Scoring Rubric": "I.3 Benchmark scoring rubric",
    "Dimension 1. Domestic Payment Interoperability": "Dimension 1. Domestic payment interoperability",
    "Dimension 2. Identity and Onboarding": "Dimension 2. Identity and onboarding",
    "Dimension 3. Functional Access and Usability": "Dimension 3. Functional access and usability",
    "Dimension 4. Remittance Integration": "Dimension 4. Remittance integration",
    "Dimension 5. Fraud Protection and Recourse": "Dimension 5. Fraud protection and recourse",
    "Dimension 6. Gender and Vulnerability Sensitivity": "Dimension 6. Gender and vulnerability sensitivity",
    "Dimension 7. Governance Coordination": "Dimension 7. Governance coordination",
    "I.4 Benchmark Scorecard": "I.4 Benchmark scorecard",
    "I.5 Country Notes": "I.5 Country notes",
    "I.6 Regional Pattern Matrix": "I.6 Regional pattern matrix",
    "I.7 Source Matrix": "I.7 Source matrix",
}

ACRONYMS = {
    "iadt": "IADT",
    "iaer": "IAER",
    "iaff": "IAFF",
    "iat": "IAT",
    "ibpd": "IBPD",
    "icdp": "ICDP",
    "icpf": "ICPF",
    "iedf": "IEDF",
    "ieh": "IEH",
    "ietr": "IETR",
    "ipcs": "IPCS",
    "iuof": "IUOF",
    "iurd": "IURD",
    "ivs": "IVS",
    "oqi": "OQI",
}


def text_of(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t"))).strip()


def clean_cell_value(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return f"{value:.6g}"
    return str(value).replace("\u2011", "-").strip()


def trim_rows(rows: Iterable[Sequence[object]]) -> list[list[str]]:
    cleaned = [[clean_cell_value(cell) for cell in row] for row in rows]
    cleaned = [row for row in cleaned if any(cell for cell in row)]
    if not cleaned:
        return []
    width = max(max((index + 1 for index, cell in enumerate(row) if cell), default=0) for row in cleaned)
    return [(row + [""] * width)[:width] for row in cleaned]


def read_workbook(path: Path) -> dict[str, list[list[str]]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    output: dict[str, list[list[str]]] = {}
    for sheet in workbook.worksheets:
        if sheet.sheet_state != "visible":
            continue
        dimensions = sheet.calculate_dimension(force=True)
        _, _, max_col, max_row = range_boundaries(dimensions)
        rows = trim_rows(
            sheet.iter_rows(
                min_row=1,
                max_row=min(max_row, 5000),
                max_col=min(max_col, 40),
                values_only=True,
            )
        )
        if rows:
            output[sheet.title] = rows
    workbook.close()
    return output


def set_style_fonts(style, size: float, *, bold: bool, italic: bool, underline: bool) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.underline = underline
    run_properties = style.element.get_or_add_rPr()
    fonts = run_properties.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), "Times New Roman")


def set_style_outline_body_text(style) -> None:
    paragraph_properties = style.element.get_or_add_pPr()
    outline = paragraph_properties.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        paragraph_properties.append(outline)
    outline.set(qn("w:val"), "9")


def ensure_appendix_styles(document: Document) -> None:
    specifications = {
        APPENDIX_TITLE_STYLE: (12, True, False, True, 12, 6),
        APPENDIX_SECTION_STYLE: (11, True, False, True, 10, 4),
        APPENDIX_SUBHEADING_STYLE: (11, True, False, False, 8, 3),
        APPENDIX_MINOR_STYLE: (11, True, True, False, 6, 2),
        APPENDIX_NOTE_STYLE: (9, False, True, False, 2, 6),
    }
    for name, (size, bold, italic, underline, before, after) in specifications.items():
        try:
            style = document.styles[name]
        except KeyError:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
        set_style_fonts(style, size, bold=bold, italic=italic, underline=underline)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        set_style_outline_body_text(style)


def get_or_add_run_properties(run_element):
    properties = run_element.find(qn("w:rPr"))
    if properties is None:
        properties = OxmlElement("w:rPr")
        run_element.insert(0, properties)
    return properties


def set_toggle(properties, tag: str, value: bool | None) -> None:
    for element in list(properties.findall(qn(f"w:{tag}"))):
        properties.remove(element)
    if value is not None:
        element = OxmlElement(f"w:{tag}")
        element.set(qn("w:val"), "1" if value else "0")
        properties.append(element)


def set_run_format(
    run_element,
    *,
    size: float,
    color: str = TEXT_HEX,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
) -> None:
    properties = get_or_add_run_properties(run_element)
    for hidden_tag in ("vanish", "webHidden"):
        for hidden in list(properties.findall(qn(f"w:{hidden_tag}"))):
            properties.remove(hidden)
    fonts = properties.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), "Times New Roman")
    for tag in ("sz", "szCs"):
        for old in list(properties.findall(qn(f"w:{tag}"))):
            properties.remove(old)
        element = OxmlElement(f"w:{tag}")
        element.set(qn("w:val"), str(round(size * 2)))
        properties.append(element)
    for old in list(properties.findall(qn("w:color"))):
        properties.remove(old)
    color_element = OxmlElement("w:color")
    color_element.set(qn("w:val"), color)
    properties.append(color_element)
    if bold is not None:
        set_toggle(properties, "b", bold)
        set_toggle(properties, "bCs", bold)
    if italic is not None:
        set_toggle(properties, "i", italic)
        set_toggle(properties, "iCs", italic)
    if underline is not None:
        for old in list(properties.findall(qn("w:u"))):
            properties.remove(old)
        if underline:
            element = OxmlElement("w:u")
            element.set(qn("w:val"), "single")
            properties.append(element)


def format_paragraph_runs(
    paragraph: Paragraph,
    *,
    size: float,
    color: str = TEXT_HEX,
    bold: bool | None = None,
    italic: bool | None = None,
    underline: bool | None = None,
) -> None:
    for run_element in paragraph._p.iter(qn("w:r")):
        if run_element.find(".//" + qn("w:drawing")) is not None:
            continue
        set_run_format(
            run_element,
            size=size,
            color=color,
            bold=bold,
            italic=italic,
            underline=underline,
        )


def set_paragraph_body_outline(paragraph: Paragraph) -> None:
    properties = paragraph._p.get_or_add_pPr()
    for old in list(properties.findall(qn("w:outlineLvl"))):
        properties.remove(old)
    outline = OxmlElement("w:outlineLvl")
    outline.set(qn("w:val"), "9")
    properties.append(outline)


def replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.text == text:
        return
    properties = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not properties:
            paragraph._p.remove(child)
    paragraph.add_run(text)


def normalize_heading_text(text: str) -> str:
    text = HEADING_REPLACEMENTS.get(text, text)
    match = re.match(r"^(A\.[1-5])\.\s+(.+?)[.]?$", text)
    if match:
        text = f"{match.group(1)} {match.group(2).rstrip('.')}"
    return text


def sentence_case(text: str) -> str:
    if not text:
        return text
    return text[:1].upper() + text[1:].lower()


def humanize_sheet_name(name: str) -> str:
    if name in SHEET_TITLES:
        return SHEET_TITLES[name]
    model = re.fullmatch(r"([MH]\d+)_k(\d+)", name, flags=re.IGNORECASE)
    if model:
        return f"Model {model.group(1).upper()}, k={model.group(2)}"
    return sentence_case(name.replace("_", " ").strip())


def clean_table_caption(text: str) -> str:
    text = re.sub(r"^Table ([HI])(\d+)\.", r"Table \1.\2.", text)
    match = re.match(r"^Table ([FG])\.(\d+)\.\s*(.+)$", text)
    if not match:
        return text.replace(" - Panel ", ", Panel ")
    appendix, number_text, remainder = match.groups()
    number = int(number_text)
    expected_title = (F_TABLE_TITLES if appendix == "F" else G_TABLE_TITLES).get(number)
    if expected_title and (
        remainder == expected_title
        or remainder.startswith(expected_title + ":")
        or remainder.startswith(expected_title + ", Panel ")
    ):
        return text
    pieces = remainder.split(" - ")
    panel = None
    if pieces and re.fullmatch(r"Panel [A-Z]", pieces[-1]):
        panel = pieces.pop()
    title = (F_TABLE_TITLES if appendix == "F" else G_TABLE_TITLES).get(number, sentence_case(pieces[0]))
    descriptor = pieces[1] if len(pieces) > 1 else None
    result = f"Table {appendix}.{number}. {title}"
    if descriptor:
        result += f": {humanize_sheet_name(descriptor)}"
    if panel:
        result += f", {panel}"
    return result


def measure_label(raw: str) -> str:
    words = raw.split()
    if not words:
        return raw
    acronym = ACRONYMS.get(words[0].lower(), words[0].upper())
    suffix = ""
    lowered = [word.lower() for word in words[1:]]
    if lowered == ["score", "01"]:
        suffix = " score (0-1 scale)"
    elif lowered == ["score"]:
        suffix = " score"
    return acronym + suffix


def clean_figure_caption(text: str) -> str:
    match = re.match(r"^Figure F\.(\d+)\.\s*(.+)$", text)
    if match:
        number = int(match.group(1))
        raw = match.group(2)
        if raw.startswith(
            (
                "Categorical distribution of ",
                "Box plot of ",
                "Histogram of ",
                "Kernel-density estimate of ",
            )
        ):
            return text
        if number in F_FIGURE_TITLES:
            title = F_FIGURE_TITLES[number]
        elif raw.startswith("Cat Distribution ") and raw.endswith(" Cat"):
            title = f"Categorical distribution of {measure_label(raw[17:-4])}"
        elif raw.startswith("Dist Box "):
            title = f"Box plot of {measure_label(raw[9:])}"
        elif raw.startswith("Dist Hist "):
            title = f"Histogram of {measure_label(raw[10:])}"
        elif raw.startswith("Dist Kdensity "):
            title = f"Kernel-density estimate of {measure_label(raw[14:])}"
        else:
            title = sentence_case(raw)
        return f"Figure F.{number}. {title}"
    match = re.match(r"^Figure G\.(\d+)\.\s*(.+)$", text)
    if match:
        number = int(match.group(1))
        return f"Figure G.{number}. {G_FIGURE_TITLES.get(number, sentence_case(match.group(2)))}"
    return text


def clean_caption(text: str) -> str:
    if text.startswith("Table "):
        text = clean_table_caption(text)
        text = HEADING_REPLACEMENTS.get(text, text)
    elif text.startswith("Figure "):
        text = clean_figure_caption(text)
    return text


def is_title_like(text: str) -> bool:
    if not (2 <= len(text.split()) <= 13) or len(text) > 105:
        return False
    if ":" in text:
        return False
    if text.startswith(('"', "'", "(", "☐")) or text.endswith((".", "?", "!", ":")):
        return False
    if "@" in text or "www." in text.lower():
        return False
    tokens = re.findall(r"[A-Za-zÀ-ÖØ-öø-ÿ][A-Za-zÀ-ÖØ-öø-ÿ/-]*", text)
    if not tokens:
        return False
    connectors = {"and", "or", "of", "the", "to", "for", "in", "with", "by", "vs"}
    substantive = [token for token in tokens if token.lower() not in connectors]
    title_tokens = [token for token in substantive if token[0].isupper() or token.isupper()]
    return bool(substantive) and len(title_tokens) / len(substantive) >= 0.70


def classify_paragraph(text: str, style_name: str, appendix: str) -> str | None:
    if text == "Appendix":
        return None
    if re.match(r"^Appendix [A-I]\.", text):
        return APPENDIX_TITLE_STYLE
    if re.match(r"^(?:Table|Figure) [A-I](?:\.|\d)", text):
        return "Caption"
    if text.lower() == "appendix note":
        return APPENDIX_NOTE_STYLE
    if re.match(r"^[A-I]\.\d+(?:\.\d+)?(?:\.|\s)", text):
        return APPENDIX_SECTION_STYLE
    if appendix == "A":
        if re.match(r"^\d+\.\d+\s", text):
            return APPENDIX_SUBHEADING_STYLE
        if re.match(r"^\d+\.\s", text):
            colon = text.find(":")
            if colon < 0 or len(text[colon + 1 :].strip()) <= 10:
                return APPENDIX_SUBHEADING_STYLE
        if re.match(r"^Section [A-Z]\.", text):
            return APPENDIX_SUBHEADING_STYLE
        if re.search(r"\(\d+(?:[–-]\d+)? minutes\)$", text):
            return APPENDIX_MINOR_STYLE
        if is_title_like(text):
            return APPENDIX_MINOR_STYLE
    if style_name.startswith("Heading"):
        try:
            level = int(style_name.split()[-1])
        except ValueError:
            level = 3
        return APPENDIX_SUBHEADING_STYLE if level <= 2 else APPENDIX_MINOR_STYLE
    return None


def bold_leading_label(paragraph: Paragraph) -> bool:
    text = paragraph.text.strip()
    if not text or paragraph._p.find(".//" + qn("w:hyperlink")) is not None:
        return False
    if paragraph._p.find(".//" + qn("w:drawing")) is not None:
        return False
    colon = text.find(":")
    if colon < 1 or colon > 55:
        return False
    prefix = text[: colon + 1]
    remainder = text[colon + 1 :]
    properties = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not properties:
            paragraph._p.remove(child)
    leading_run = paragraph.add_run(prefix)
    leading_run.bold = True
    if remainder:
        paragraph.add_run(remainder)
    return True


def style_top_level_paragraph(paragraph: Paragraph, appendix: str) -> tuple[str, bool]:
    text = paragraph.text.strip()
    original_style = paragraph.style.name if paragraph.style is not None else "Normal"
    if text == "Appendix":
        # This is the sole appendix heading intentionally retained in the TOC.
        paragraph.style = "Heading 1"
        return appendix, False

    if re.match(r"^Appendix [A-I]\.", text):
        appendix = text.split()[1].rstrip(".")

    cleaned = normalize_heading_text(text)
    if original_style == "Caption" or re.match(r"^(?:Table|Figure) [A-I](?:\.|\d)", cleaned):
        cleaned = clean_caption(cleaned)
    replace_paragraph_text(paragraph, cleaned)

    target_style = classify_paragraph(cleaned, original_style, appendix)
    if target_style:
        paragraph.style = target_style

    set_paragraph_body_outline(paragraph)

    if target_style == APPENDIX_TITLE_STYLE:
        format_paragraph_runs(paragraph, size=12, bold=True, italic=False, underline=True)
    elif target_style == APPENDIX_SECTION_STYLE:
        format_paragraph_runs(paragraph, size=11, bold=True, italic=False, underline=True)
    elif target_style == APPENDIX_SUBHEADING_STYLE:
        format_paragraph_runs(paragraph, size=11, bold=True, italic=False, underline=False)
    elif target_style == APPENDIX_MINOR_STYLE:
        format_paragraph_runs(paragraph, size=11, bold=True, italic=True, underline=False)
    elif target_style == APPENDIX_NOTE_STYLE:
        format_paragraph_runs(paragraph, size=9, bold=False, italic=True, underline=False)
    elif target_style == "Caption" or original_style == "Caption":
        paragraph.style = "Caption"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(7)
        paragraph.paragraph_format.keep_with_next = True
        format_paragraph_runs(paragraph, size=11, bold=True, italic=True, underline=False)
    elif paragraph._p.find(".//" + qn("w:drawing")) is not None:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
    else:
        if appendix == "A":
            bold_leading_label(paragraph)
        format_paragraph_runs(paragraph, size=11)
    return appendix, bool(target_style)


def set_cell_margins(cell, top: int = 70, start: int = 85, bottom: int = 70, end: int = 85) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_width(cell, width_twips: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        properties.insert(0, width)
    width.set(qn("w:w"), str(width_twips))
    width.set(qn("w:type"), "dxa")


def set_fixed_table_geometry(table: Table, widths: Sequence[float], total_width: float = 6.45) -> None:
    total_twips = round(total_width * 1440)
    column_twips = [round(width * 1440) for width in widths]
    column_twips[-1] += total_twips - sum(column_twips)
    properties = table._tbl.tblPr
    table_width = properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.insert(0, table_width)
    table_width.set(qn("w:w"), str(total_twips))
    table_width.set(qn("w:type"), "dxa")
    table_indent = properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        properties.append(table_indent)
    table_indent.set(qn("w:w"), "0")
    table_indent.set(qn("w:type"), "dxa")
    for grid_column, width_twips in zip(table._tbl.tblGrid.findall(qn("w:gridCol")), column_twips):
        grid_column.set(qn("w:w"), str(width_twips))
    for row in table.rows:
        for index, cell in enumerate(row.cells[: len(column_twips)]):
            set_cell_width(cell, column_twips[index])


def repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    for old in list(properties.findall(qn("w:tblHeader"))):
        properties.remove(old)
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def clear_row_height_and_prevent_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    for old in list(properties.findall(qn("w:trHeight"))):
        properties.remove(old)
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def format_table(table: Table) -> None:
    if not table.rows:
        return
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    column_count = max((len(row.cells) for row in table.rows), default=1)
    font_size = 8.0 if column_count >= 6 else 9.0
    header_size = 8.5 if column_count >= 6 else 9.5
    for row_index, row in enumerate(table.rows):
        clear_row_height_and_prevent_split(row)
        if row_index == 0:
            repeat_header(row)
        for cell in row.cells:
            properties = cell._tc.get_or_add_tcPr()
            for no_wrap in list(properties.findall(qn("w:noWrap"))):
                properties.remove(no_wrap)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade_cell(cell, TEAL if row_index == 0 else (PALE_GRAY if row_index % 2 == 0 else "FFFFFF"))
            for paragraph in cell.paragraphs:
                paragraph.style = "Normal"
                set_paragraph_body_outline(paragraph)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.keep_with_next = False
                if row_index == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    format_paragraph_runs(
                        paragraph,
                        size=header_size,
                        color=WHITE_HEX,
                        bold=True,
                        italic=False,
                        underline=False,
                    )
                else:
                    format_paragraph_runs(paragraph, size=font_size, color=TEXT_HEX)


def calculate_widths(rows: list[list[str]], total_width: float = 6.45) -> list[float]:
    column_count = len(rows[0])
    lengths = []
    for column in range(column_count):
        maximum = max((len(row[column]) for row in rows), default=5)
        lengths.append(max(8, min(maximum, 42)))
    lengths[0] = max(lengths[0], 16)
    widths = [max(0.62, total_width * length / sum(lengths)) for length in lengths]
    scale = total_width / sum(widths)
    return [width * scale for width in widths]


def make_native_table(document: Document, rows: list[list[str]]) -> Table:
    if not rows:
        raise ValueError("Cannot build a Word table from an empty worksheet")
    column_count = len(rows[0])
    table = document.add_table(rows=len(rows), cols=column_count)
    for row_index, source_row in enumerate(rows):
        for column_index, value in enumerate(source_row):
            table.cell(row_index, column_index).text = value
    set_fixed_table_geometry(table, calculate_widths(rows))
    format_table(table)
    return table


def repair_architecture_tables(document: Document, appendix_start: int) -> dict[str, dict[str, int]]:
    workbook = read_workbook(ARCHITECTURE_WORKBOOK)
    required = {
        "Table C.1. Current quantitative index architecture": "Index Architecture",
        "Table C.2. Survey items, constructed components, and current scoring rules": "Item Mapping",
    }
    body = document._element.body
    children = list(body)
    repaired: dict[str, dict[str, int]] = {}
    last_caption = ""
    for child in children[appendix_start:]:
        if child.tag == qn("w:p"):
            current_text = text_of(child)
            last_caption = current_text if current_text.startswith("Table ") else last_caption
            continue
        if child.tag != qn("w:tbl"):
            continue
        if last_caption not in required:
            last_caption = ""
            continue
        sheet_name = required[last_caption]
        rows = workbook.get(sheet_name)
        if not rows:
            raise RuntimeError(f"Worksheet {sheet_name!r} was not populated")
        old_rows = child.findall(qn("w:tr"))
        old_columns = max((len(row.findall(qn("w:tc"))) for row in old_rows), default=0)
        replacement = make_native_table(document, rows)
        child.addprevious(replacement._tbl)
        body.remove(child)
        repaired[last_caption] = {
            "old_rows": len(old_rows),
            "old_columns": old_columns,
            "new_rows": len(rows),
            "new_columns": len(rows[0]),
        }
        last_caption = ""
    missing = set(required) - set(repaired)
    if missing:
        raise RuntimeError(f"Could not locate table(s) for replacement: {sorted(missing)}")
    return repaired


def find_appendix_start(document: Document) -> int:
    for index, child in enumerate(document._element.body):
        if child.tag == qn("w:p") and text_of(child).lower() == "appendix":
            return index
    raise RuntimeError("The main Appendix section title was not found")


def enable_field_updates(document: Document) -> None:
    settings = document.settings.element
    for old in list(settings.findall(qn("w:updateFields"))):
        settings.remove(old)
    element = OxmlElement("w:updateFields")
    element.set(qn("w:val"), "true")
    settings.append(element)


def audit_final_document(document: Document, appendix_start: int) -> dict[str, object]:
    heading_styles = []
    internal_outline_levels = []
    linked_images = 0
    embedded_images = 0
    alt_chunks = 0
    tables = 0
    sparse_tables = []
    current_caption = ""
    children = list(document._element.body)
    for child in children[appendix_start:]:
        for blip in child.iter(qn("a:blip")):
            linked_images += bool(blip.get(qn("r:link")))
            embedded_images += bool(blip.get(qn("r:embed")))
        alt_chunks += len(child.findall(".//" + qn("w:altChunk")))
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document._body)
            if paragraph.text == "Appendix":
                continue
            if paragraph.style is not None and paragraph.style.name.startswith("Heading"):
                heading_styles.append(paragraph.text)
            outline = child.find(qn("w:pPr") + "/" + qn("w:outlineLvl"))
            if outline is not None and outline.get(qn("w:val")) != "9":
                internal_outline_levels.append(paragraph.text)
            if paragraph.text.startswith("Table "):
                current_caption = paragraph.text
        elif child.tag == qn("w:tbl"):
            tables += 1
            row_elements = child.findall(qn("w:tr"))
            cell_values = [[text_of(cell) for cell in row.findall(qn("w:tc"))] for row in row_elements]
            body_values = [value for row in cell_values[1:] for value in row]
            if len(row_elements) <= 1 or not any(body_values):
                sparse_tables.append({"caption": current_caption, "rows": len(row_elements)})
            current_caption = ""
    return {
        "appendix_internal_heading_style_count": len(heading_styles),
        "appendix_internal_heading_styles": heading_styles,
        "appendix_internal_nonbody_outline_count": len(internal_outline_levels),
        "appendix_internal_nonbody_outline": internal_outline_levels,
        "appendix_table_count": tables,
        "appendix_sparse_table_count": len(sparse_tables),
        "appendix_sparse_tables": sparse_tables,
        "appendix_embedded_image_count": embedded_images,
        "appendix_linked_image_count": linked_images,
        "appendix_altchunk_count": alt_chunks,
    }


def polish(source: Path, output: Path) -> dict[str, object]:
    document = Document(source)
    ensure_appendix_styles(document)
    appendix_start = find_appendix_start(document)
    repaired = repair_architecture_tables(document, appendix_start)

    body = document._element.body
    appendix = "Front"
    styled_headings = 0
    caption_edits = 0
    for child in list(body)[appendix_start:]:
        if child.tag == qn("w:p"):
            paragraph = Paragraph(child, document._body)
            before = paragraph.text
            appendix, styled = style_top_level_paragraph(paragraph, appendix)
            styled_headings += int(styled)
            caption_edits += int(before != paragraph.text and before.startswith(("Table ", "Figure ")))
        elif child.tag == qn("w:tbl"):
            format_table(Table(child, document._body))

    enable_field_updates(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)

    verification_document = Document(output)
    verification_start = find_appendix_start(verification_document)
    audit = audit_final_document(verification_document, verification_start)
    audit.update(
        {
            "source": str(source.resolve()),
            "output": str(output.resolve()),
            "architecture_tables_repaired": repaired,
            "appendix_display_headings_styled": styled_headings,
            "captions_proofread": caption_edits,
        }
    )
    return audit


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--report", type=Path)
    return parser.parse_args()


if __name__ == "__main__":
    arguments = parse_args()
    result = polish(arguments.source, arguments.output)
    if arguments.report:
        arguments.report.parent.mkdir(parents=True, exist_ok=True)
        arguments.report.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(result, ensure_ascii=False, indent=2))
