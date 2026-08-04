"""Assemble the frozen quantitative online appendices into the working paper.

This script formats and embeds existing aggregate outputs. It does not estimate,
alter, or reinterpret the project's indices, regressions, interactions, LCA
models, posterior assignments, or segment definitions.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import re
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.utils.cell import range_boundaries
from PIL import Image


REPO_ROOT = Path(__file__).resolve().parents[2]
APPENDIX_ROOT = REPO_ROOT / "3 Research Paper Online Appendix"

TEAL = "0F4C5C"
MID_TEAL = "2F7A8A"
LIGHT_TEAL = "EAF3F5"
PALE_GRAY = "F4F6F7"
GRID = "B7CDD2"
TEXT = RGBColor(31, 41, 51)

APPENDIX_TITLE_STYLE = "Appendix Title"
APPENDIX_SECTION_STYLE = "Appendix Section"


DESCRIPTIVE_TITLES = {
    1: "Age distribution",
    2: "City of residence",
    3: "Years living in Colombia",
    4: "Educational attainment",
    5: "Current occupation",
    6: "Socioeconomic Vulnerability Index (IVS)",
    7: "Socioeconomic vulnerability by city",
    8: "Socioeconomic vulnerability by age group",
    9: "Mobile-phone and internet access",
    10: "Mobile-data stability",
    11: "Ability to send and receive digital money",
    12: "Telecommunications access and digital self-efficacy",
    13: "Use of QR payments",
    14: "Practical digital competence by age",
    15: "Practical digital competence by education",
    16: "Types of financial accounts owned",
    17: "Formal financial access by migration tenure",
    18: "Formal financial access by education",
    19: "Financial use and operability (IUOF)",
    20: "Financial operability and formal financial access",
    21: "Awareness and use of payment rails",
    22: "Onboarding quality by formal financial access",
    23: "Onboarding quality by practical digital competence",
    24: "Channels used to receive remittances",
    25: "Transactional experience in remittances and payments (IETR)",
    26: "Digital-remittance intensity by formal-channel use",
    27: "Perceived safety and fraud exposure",
    28: "Fraud exposure and recourse harm by onboarding quality",
    29: "Fraud exposure and recourse harm by practical digital competence",
    30: "Resolution satisfaction after payment or remittance problems",
    31: "Economic autonomy in remittances (IAER)",
    32: "Economic autonomy by digital-remittance intensity",
    33: "Fraud exposure and recourse harm by economic autonomy",
    34: "Trust and financial climate by digital-remittance intensity",
    35: "Five most frequently reported barriers",
    36: "Five most frequently requested enabling supports",
    37: "Priority training topics",
    38: "Perceived digitalization barriers by socioeconomic vulnerability",
    39: "Perceived digitalization barriers by telecommunications access",
}

REGRESSION_TITLES = {
    40: "Coefficient plot: digital-remittance and payment intensity (IURD)",
    41: "Average marginal effects: formal remittance-channel use",
    42: "Coefficient plot: financial use and operability (IUOF)",
    43: "Coefficient plot: transactional experience (IETR)",
    44: "Coefficient plot: onboarding quality (OQI)",
    45: "Coefficient plot: prevention and safe conduct (IPCS)",
    46: "Coefficient plot: fraud exposure and recourse harm (IEDF)",
    47: "Coefficient plot: economic autonomy in remittances (IAER)",
    48: "Coefficient plot: trust and financial climate (ICPF)",
    49: "Interaction: socioeconomic vulnerability and telecommunications access",
    50: "Interaction: migration tenure and formal financial access",
    51: "Interaction: practical digital competence and onboarding quality",
    52: "Interaction: fraud harm and trust/financial climate",
}

REGRESSION_TABLE_TITLES = {
    "Table_E1_IURD.txt": "Digital-remittance and payment intensity (IURD)",
    "Table_E2_Formal_Remittance.txt": "Formal remittance-channel use",
    "Table_E3_IUOF.txt": "Financial use and operability (IUOF)",
    "Table_E4_IETR.txt": "Transactional experience in remittances and payments (IETR)",
    "Table_E5_OQI.txt": "Onboarding quality (OQI)",
    "Table_E6_IPCS.txt": "Prevention and safe conduct (IPCS)",
    "Table_E7_IEDF.txt": "Fraud exposure and recourse harm (IEDF)",
    "Table_E8_IAER.txt": "Economic autonomy in remittances (IAER)",
    "Table_E9_ICPF.txt": "Trust and financial climate (ICPF)",
    "Table_E10_Interactions.txt": "Interaction specifications",
    "Table_E11_Supplementary_Items.txt": "Supplementary item-level models",
}

DISPLAY_LABELS = {
    "2.age_cat": "Age 30-44",
    "3.age_cat": "Age 45-59",
    "4.age_cat": "Age 60+",
    "2.q3": "Cali",
    "3.q3": "Medellin",
    "4.q3": "Soacha",
    "ivs_score": "Socioeconomic vulnerability (IVS)",
    "iat_score": "Telecommunications access (IAT)",
    "iadt_score": "Digital transactional self-efficacy (IADT)",
    "icdp_score": "Practical digital competence (ICDP)",
    "iaff_score": "Formal financial access (IAFF)",
    "iuof_score_01": "Financial use and operability (IUOF)",
    "oqi_score_01": "Onboarding quality (OQI)",
    "iurd_score_01": "Digital-remittance intensity (IURD)",
    "ietr_score_01": "Transactional experience (IETR)",
    "IPCS": "Prevention and safe conduct (IPCS)",
    "IEDF": "Fraud exposure and recourse harm (IEDF)",
    "IAER": "Economic autonomy in remittances (IAER)",
    "ICPF": "Trust and financial climate (ICPF)",
    "IEH": "Enabling environment (IEH)",
    "IBPD": "Perceived barriers to digitalization (IBPD)",
    "years_in_col": "Years living in Colombia",
    "formal_remittance": "Formal remittance-channel use",
    "Constant": "Constant",
    "Observations": "Observations",
    "R-squared": "R-squared",
    "City FE": "City fixed effects",
    "Robust SE": "Robust standard errors",
}


def clean_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return f"{value:.6g}"
    text = str(value).replace("\u2011", "-").strip()
    if "Ã" in text or "Â" in text:
        try:
            text = text.encode("cp1252").decode("utf-8")
        except (UnicodeEncodeError, UnicodeDecodeError):
            pass
    return DISPLAY_LABELS.get(text, text)


def trim_rows(rows: Iterable[Sequence[object]]) -> list[list[str]]:
    cleaned = [[clean_text(cell) for cell in row] for row in rows]
    cleaned = [row for row in cleaned if any(cell != "" for cell in row)]
    if not cleaned:
        return []
    width = max(max((idx + 1 for idx, cell in enumerate(row) if cell != ""), default=0) for row in cleaned)
    if width == 0:
        return []
    return [(row + [""] * width)[:width] for row in cleaned]


def read_delimited(path: Path, delimiter: str) -> list[list[str]]:
    with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as stream:
        rows = list(csv.reader(stream, delimiter=delimiter))
    rows = trim_rows(rows)
    filtered: list[list[str]] = []
    for row in rows:
        if row and row[0] in {"formal_remittance", "Structurally_constrained_low_int", "Operationally_active_but_high_ri", "Mainstream_formal_digital_users", "Integrated_and_protected_digital"}:
            if all(cell in {"", ".", "(.)"} for cell in row[1:]):
                continue
        filtered.append([DISPLAY_LABELS.get(cell, cell) if idx == 0 else cell for idx, cell in enumerate(row)])
    return filtered


def read_workbook(path: Path) -> list[tuple[str, list[list[str]]]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    output: list[tuple[str, list[list[str]]]] = []
    for sheet in workbook.worksheets:
        if sheet.sheet_state != "visible":
            continue
        # Artifact-tool workbooks can omit the cached worksheet dimensions.
        # In read-only mode openpyxl then reports max_row/max_column as None,
        # even though the worksheet contains data. Force a dimension scan so
        # every populated cell is embedded in the native Word table.
        dimensions = sheet.calculate_dimension(force=True)
        _, _, dimension_max_col, dimension_max_row = range_boundaries(dimensions)
        max_row = min(dimension_max_row, 5000)
        max_col = min(dimension_max_col, 40)
        rows = trim_rows(sheet.iter_rows(min_row=1, max_row=max_row, max_col=max_col, values_only=True))
        if rows:
            output.append((sheet.title, rows))
    workbook.close()
    return output


def shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 70, start: int = 85, bottom: int = 70, end: int = 85) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{tag}"))
        if element is None:
            element = OxmlElement(f"w:{tag}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_fixed_table_geometry(table, widths: Sequence[float], total_width: float) -> list[int]:
    """Write a consistent fixed-width table grid in twentieths of a point."""
    total_twips = round(total_width * 1440)
    column_twips = [round(width * 1440) for width in widths]
    column_twips[-1] += total_twips - sum(column_twips)

    properties = table._tbl.tblPr
    table_width = properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.insert(0, table_width)
    table_width.set(qn("w:w"), str(total_twips))
    table_width.set(qn("w:type"), "dxa")

    table_indent = properties.first_child_found_in("w:tblInd")
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        properties.append(table_indent)
    table_indent.set(qn("w:w"), "0")
    table_indent.set(qn("w:type"), "dxa")

    grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
    for grid_column, width_twips in zip(grid_columns, column_twips):
        grid_column.set(qn("w:w"), str(width_twips))
    return column_twips


def set_cell_width(cell, width_twips: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    cell_width = properties.first_child_found_in("w:tcW")
    if cell_width is None:
        cell_width = OxmlElement("w:tcW")
        properties.insert(0, cell_width)
    cell_width.set(qn("w:w"), str(width_twips))
    cell_width.set(qn("w:type"), "dxa")


def repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def set_alt_text(shape, description: str) -> None:
    doc_properties = shape._inline.docPr
    doc_properties.set("descr", description)
    doc_properties.set("title", description)


def ensure_appendix_styles(document: Document) -> None:
    """Create appendix display styles that are intentionally excluded from the TOC."""
    specifications = {
        APPENDIX_TITLE_STYLE: {"size": 12, "bold": True, "underline": True, "before": 12, "after": 6},
        APPENDIX_SECTION_STYLE: {"size": 11, "bold": True, "underline": True, "before": 10, "after": 4},
    }
    for name, specification in specifications.items():
        try:
            style = document.styles[name]
        except KeyError:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(specification["size"])
        style.font.bold = specification["bold"]
        style.font.underline = specification["underline"]
        style.paragraph_format.space_before = Pt(specification["before"])
        style.paragraph_format.space_after = Pt(specification["after"])
        style.paragraph_format.keep_with_next = True
        paragraph_properties = style.element.get_or_add_pPr()
        outline = paragraph_properties.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            paragraph_properties.append(outline)
        # Word uses 9 for Body Text. The manuscript TOC field includes outline
        # levels, so this explicit value prevents appendix internals from being
        # pulled into the TOC during a field refresh.
        outline.set(qn("w:val"), "9")


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(7)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT


def add_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(80, 90, 98)


def page_break(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def normalized_table_panels(rows: list[list[str]], max_columns: int = 8) -> list[tuple[str | None, list[list[str]]]]:
    if not rows:
        return []
    width = max(len(row) for row in rows)
    rows = [(row + [""] * width)[:width] for row in rows]
    if width <= max_columns:
        return [(None, rows)]
    stub_columns = 1 if "correlation" in " ".join(rows[0]).lower() else min(2, width - 1)
    data_columns_per_panel = max_columns - stub_columns
    panels: list[tuple[str | None, list[list[str]]]] = []
    panel_number = 0
    for start in range(stub_columns, width, data_columns_per_panel):
        panel_number += 1
        indices = list(range(stub_columns)) + list(range(start, min(start + data_columns_per_panel, width)))
        panel_rows = [[row[index] for index in indices] for row in rows]
        panels.append((f"Panel {chr(64 + panel_number)}", panel_rows))
    return panels


def add_word_table(document: Document, rows: list[list[str]], caption: str, note: str | None = None) -> None:
    panels = normalized_table_panels(rows)
    for panel_index, (panel_name, panel_rows) in enumerate(panels):
        if panel_index > 0:
            page_break(document)
        full_caption = caption if panel_name is None else f"{caption} - {panel_name}"
        add_caption(document, full_caption)
        column_count = len(panel_rows[0])
        table = document.add_table(rows=1, cols=column_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.style = "Table Grid"
        total_width = 6.45
        header_lengths = [max(5, len(panel_rows[0][index])) for index in range(column_count)]
        content_lengths = [
            max(header_lengths[index], min(45, max((len(row[index]) for row in panel_rows[1:]), default=5)))
            for index in range(column_count)
        ]
        if column_count > 1:
            content_lengths[0] = max(content_lengths[0], 12)
        length_sum = sum(content_lengths)
        widths = [max(0.62, total_width * length / length_sum) for length in content_lengths]
        scale = total_width / sum(widths)
        widths = [width * scale for width in widths]
        column_twips = set_fixed_table_geometry(table, widths, total_width)

        for column_index, value in enumerate(panel_rows[0]):
            cell = table.rows[0].cells[column_index]
            set_cell_width(cell, column_twips[column_index])
            cell.text = value
            shade_cell(cell, TEAL)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(7.5)
                run.font.color.rgb = RGBColor(255, 255, 255)
        repeat_header(table.rows[0])
        prevent_row_split(table.rows[0])

        for row_index, source_row in enumerate(panel_rows[1:], start=1):
            cells = table.add_row().cells
            for column_index, value in enumerate(source_row):
                cell = cells[column_index]
                set_cell_width(cell, column_twips[column_index])
                cell.text = value
                set_cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if row_index % 2 == 0:
                    shade_cell(cell, PALE_GRAY)
                paragraph = cell.paragraphs[0]
                is_numeric = bool(re.fullmatch(r"[-+]?\(?\d[\d.,]*\)?\**", value.strip()))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if is_numeric else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(7.2)
                    run.font.color.rgb = TEXT
                    if value.startswith("(") and value.endswith(")"):
                        run.italic = True
                        run.font.color.rgb = RGBColor(90, 98, 104)
            prevent_row_split(table.rows[row_index])
        if note and panel_index == len(panels) - 1:
            add_note(document, note)


def add_figure(document: Document, path: Path, caption: str, source_note: str | None = None) -> None:
    with Image.open(path) as image:
        pixel_width, pixel_height = image.size
    max_width = 6.3
    max_height = 7.65
    aspect = pixel_width / pixel_height if pixel_height else 1
    width = max_width
    height = width / aspect
    if height > max_height:
        height = max_height
        width = height * aspect
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    shape = paragraph.add_run().add_picture(str(path), width=Inches(width), height=Inches(height))
    set_alt_text(shape, caption)
    add_caption(document, caption)
    if source_note:
        add_note(document, source_note)


def descriptive_figure_title(path: Path) -> str:
    match = re.match(r"fig(\d+)", path.stem, flags=re.IGNORECASE)
    number = int(match.group(1)) if match else 0
    title = DESCRIPTIVE_TITLES.get(number, path.stem.replace("_", " ").title())
    suffix = path.stem.lower()
    if "porcentage" in suffix:
        title += " - percentage display"
    elif "density" in suffix:
        title += " - density display"
    elif "distrib" in suffix and number in {6, 19, 25, 31}:
        title += " - distribution display"
    return title


def generic_title(path: Path) -> str:
    name = re.sub(r"^(Table|Figure)_[A-Z]\d+[A-Za-z]?_?", "", path.stem, flags=re.IGNORECASE)
    name = re.sub(r"^figC\d+[A-Za-z]?_?", "", name, flags=re.IGNORECASE)
    return name.replace("_", " ").strip().title()


def add_xlsx(document: Document, path: Path, label: str, title: str) -> None:
    sheets = read_workbook(path)
    for sheet_index, (sheet_name, rows) in enumerate(sheets):
        if sheet_index > 0:
            page_break(document)
        sheet_suffix = "" if len(sheets) == 1 else f" - {sheet_name}"
        add_word_table(document, rows, f"Table {label}. {title}{sheet_suffix}")


def add_major_heading(document: Document, text: str) -> None:
    page_break(document)
    paragraph = document.add_paragraph(text, style=APPENDIX_TITLE_STYLE)
    paragraph.paragraph_format.keep_with_next = True


def add_subheading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text, style=APPENDIX_SECTION_STYLE)
    paragraph.paragraph_format.keep_with_next = True


def add_intro(document: Document, paragraphs: Sequence[str]) -> None:
    for text in paragraphs:
        paragraph = document.add_paragraph(text)
        paragraph.paragraph_format.space_after = Pt(6)


def add_appendix_c(document: Document) -> None:
    add_major_heading(document, "Appendix C. Quantitative variable construction and index architecture")
    add_intro(document, [
        "This appendix documents the quantitative measurement architecture exactly as implemented in the current Stata data-preparation workflow. It is intended to make the existing index definitions, item mappings, aggregation rules, scales, and descriptive thresholds auditable.",
        "Analytical freeze. The tables below organize the current implementation and estimates; they do not correct or replace any index, regression, interaction, LCA specification, posterior assignment, or segment definition.",
    ])
    architecture_path = APPENDIX_ROOT / "C Variable Construction" / "tables" / "Table_C0_Current_Index_Architecture.xlsx"
    architecture_sheets = read_workbook(architecture_path)
    add_subheading(document, "C.1 Index architecture")
    add_word_table(document, architecture_sheets[0][1], "Table C.1. Current quantitative index architecture")
    page_break(document)
    add_subheading(document, "C.2 Item-to-index mapping")
    add_word_table(
        document,
        architecture_sheets[1][1],
        "Table C.2. Survey items, constructed components, and current scoring rules",
        "Note: This is a documentation table of the frozen current code. All row means use available nonmissing components unless the code states otherwise.",
    )
    page_break(document)
    add_subheading(document, "C.3 Index distributions and associations")
    summary = read_delimited(APPENDIX_ROOT / "C Variable Construction" / "tables" / "Table_C1_Index_Summary_Statistics.csv", ",")
    add_word_table(document, summary, "Table C.3. Summary statistics for the fifteen quantitative indices", "N=423 for each index in the frozen analytical dataset.")
    page_break(document)
    correlation = read_delimited(APPENDIX_ROOT / "C Variable Construction" / "tables" / "Table_C2_Index_Correlation_Matrix.csv", ",")
    add_word_table(document, correlation, "Table C.4. Correlation matrix across the fifteen quantitative indices")
    figures = APPENDIX_ROOT / "C Variable Construction" / "figures"
    for number, filename, title in [
        (1, "Figure_C1_Index_Distributions.png", "Distribution of the quantitative indices"),
        (2, "Figure_C2_Index_Correlation_Heatmap.png", "Correlation heatmap across the quantitative indices"),
    ]:
        path = figures / filename
        if path.exists():
            page_break(document)
            add_figure(document, path, f"Figure C.{number}. {title}", "Source: Authors' calculations from the frozen analytical dataset.")


def add_appendix_d(document: Document) -> None:
    add_major_heading(document, "Appendix D. Quantitative descriptive evidence")
    add_intro(document, [
        "This appendix preserves the extended descriptive evidence underlying the quantitative results. Percentages are unweighted and describe the achieved nonprobability sample.",
        "The tables and figures reproduce the existing analytical outputs without changing their definitions or underlying data transformations.",
    ])
    tables_dir = APPENDIX_ROOT / "D Descriptive Evidence" / "tables"
    add_subheading(document, "D.1 Extended descriptive tables")
    add_word_table(document, read_delimited(tables_dir / "Table_D1_Extended_Categorical_Distributions.csv", ","), "Table D.1. Demographic, migration, access, remittance, fraud, recourse, and support distributions")
    page_break(document)
    add_word_table(document, read_delimited(tables_dir / "Table_D2_Multiple_Response_Prevalence.csv", ","), "Table D.2. Multiple-response financial-access and enabling-support prevalence")

    add_subheading(document, "D.2 Extended descriptive figures")
    figures = sorted(
        (APPENDIX_ROOT / "D Descriptive Evidence" / "figures").glob("*.png"),
        key=lambda path: (int(re.search(r"fig(\d+)", path.name, re.IGNORECASE).group(1)), path.name.lower()),
    )
    for appendix_number, path in enumerate(figures, start=1):
        page_break(document)
        add_figure(document, path, f"Figure D.{appendix_number}. {descriptive_figure_title(path)}", "Source: Authors' calculations from the quantitative survey (N=423 unless the figure indicates an applicable subsample).")


def add_appendix_e(document: Document) -> None:
    add_major_heading(document, "Appendix E. Multivariate regression analysis")
    add_intro(document, [
        "The appendix reports the existing nested regression specifications. Models retain the original outcome definitions, covariates, estimators, robust standard errors, interactions, and analytical samples.",
        "Models A-D progressively add the current index blocks. City fixed effects and age-category controls appear as shown in each table. Binary formal-remittance models report odds ratios; continuous-index models report the existing OLS coefficients.",
    ])
    add_subheading(document, "E.1 Complete regression tables")
    tables_dir = APPENDIX_ROOT / "E Multivariate Regressions" / "tables"
    table_files = sorted(tables_dir.glob("*.txt"), key=lambda path: int(re.search(r"Table_E(\d+)", path.name).group(1)))
    for path in table_files:
        label = re.search(r"Table_(E\d+)", path.name).group(1).replace("E", "E.")
        title = REGRESSION_TABLE_TITLES.get(path.name, generic_title(path))
        page_break(document)
        rows = read_delimited(path, "\t")
        add_word_table(document, rows, f"Table {label}. {title}", "Robust standard errors are shown in parentheses. Significance markers reproduce the frozen regression output.")

    add_subheading(document, "E.2 Coefficient, marginal-effect, and interaction figures")
    figures = sorted(
        (APPENDIX_ROOT / "E Multivariate Regressions" / "figures").glob("*.png"),
        key=lambda path: int(re.search(r"fig(\d+)", path.name, re.IGNORECASE).group(1)),
    )
    for appendix_number, path in enumerate(figures, start=1):
        source_number = int(re.search(r"fig(\d+)", path.name, re.IGNORECASE).group(1))
        page_break(document)
        add_figure(document, path, f"Figure E.{appendix_number}. {REGRESSION_TITLES[source_number]}", "Source: Existing frozen regression and margins outputs.")


def add_appendix_f(document: Document) -> None:
    add_major_heading(document, "Appendix F. Latent class analysis diagnostics and model selection")
    add_intro(document, [
        "This appendix documents the existing latent class analysis workflow, including sample integrity, candidate indicators, recoding, feature sets, candidate-model comparisons, convergence records, final fit, response patterns, and posterior-classification diagnostics.",
        "No candidate model, selected model, indicator coding, class assignment, or posterior probability has been re-estimated or changed for this appendix.",
    ])
    add_subheading(document, "F.1 Diagnostic and model-selection tables")
    tables_dir = APPENDIX_ROOT / "F LCA Diagnostics" / "tables"
    table_files = sorted(tables_dir.glob("*.xlsx"), key=lambda path: int(re.search(r"Table_F(\d+)", path.name).group(1)))
    for path in table_files:
        number = int(re.search(r"Table_F(\d+)", path.name).group(1))
        page_break(document)
        add_xlsx(document, path, f"F.{number}", generic_title(path))

    add_subheading(document, "F.2 Indicator-distribution and classification-diagnostic figures")
    figure_dir = APPENDIX_ROOT / "F LCA Diagnostics" / "figures"
    figures = []
    for path in figure_dir.glob("*.png"):
        if path.name.startswith(("figC6_", "figC7_", "figC10")):
            continue
        figures.append(path)
    figures.sort(key=lambda path: path.name.lower())
    for number, path in enumerate(figures, start=1):
        page_break(document)
        add_figure(document, path, f"Figure F.{number}. {generic_title(path)}", "Source: Existing frozen LCA diagnostic output.")


def add_appendix_g(document: Document) -> None:
    add_major_heading(document, "Appendix G. Final LCA segment profiles and post-LCA validation")
    add_intro(document, [
        "This appendix reports the existing four-segment solution, class mapping, conditional response probabilities, continuous centroids, signature behaviors, demographic and financial profiles, auxiliary outcomes, membership predictors, and outcome-regression overlays.",
        "Terminology and statistical results reproduce the frozen segmentation workflow. Class-defining indicators and auxiliary profiles are presented separately in the underlying tables.",
    ])
    add_subheading(document, "G.1 Final segment tables")
    tables_dir = APPENDIX_ROOT / "G Segment Profiles" / "tables"
    paths = sorted(tables_dir.iterdir(), key=lambda path: int(re.search(r"Table_G(\d+)", path.name).group(1)))
    for path in paths:
        number = int(re.search(r"Table_G(\d+)", path.name).group(1))
        page_break(document)
        if path.suffix.lower() == ".xlsx":
            add_xlsx(document, path, f"G.{number}", generic_title(path))
        elif path.suffix.lower() == ".txt":
            add_word_table(document, read_delimited(path, "\t"), f"Table G.{number}. {generic_title(path)}")

    add_subheading(document, "G.2 Segment-profile and post-LCA figures")
    figures = sorted(
        (APPENDIX_ROOT / "G Segment Profiles" / "figures").glob("*.png"),
        key=lambda path: int(re.search(r"Figure_G(\d+)", path.name).group(1)),
    )
    for number, path in enumerate(figures, start=1):
        page_break(document)
        add_figure(document, path, f"Figure G.{number}. {generic_title(path)}", "Source: Existing frozen final-segment and post-LCA output.")


def remove_placeholder_block(document: Document):
    body = document._element.body
    children = list(body)
    start = None
    anchor = None
    for child in children:
        if child.tag != qn("w:p"):
            continue
        text = "".join(child.itertext()).strip()
        if start is None and text.startswith("Appendix F. Latent class analysis diagnostics"):
            start = child
        if text.startswith("Appendix H. Qualitative coding"):
            anchor = child
            break
    if anchor is None:
        raise RuntimeError("Appendix H anchor was not found in the source manuscript")
    if start is not None:
        deleting = False
        for child in list(body):
            if child is start:
                deleting = True
            if child is anchor:
                break
            if deleting:
                body.remove(child)
    return anchor


def enable_field_updates(document: Document) -> None:
    settings = document.settings.element
    for existing in settings.findall(qn("w:updateFields")):
        settings.remove(existing)
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def assemble(source: Path, output: Path) -> None:
    document = Document(source)
    ensure_appendix_styles(document)
    anchor = remove_placeholder_block(document)
    body = document._element.body
    before_children = list(body)
    has_final_section_properties = bool(before_children and before_children[-1].tag == qn("w:sectPr"))
    insertion_start = len(before_children) - (1 if has_final_section_properties else 0)

    add_appendix_c(document)
    add_appendix_d(document)
    add_appendix_e(document)
    add_appendix_f(document)
    add_appendix_g(document)

    after_children = list(body)
    insertion_end = len(after_children) - (1 if has_final_section_properties else 0)
    new_elements = after_children[insertion_start:insertion_end]
    for element in new_elements:
        anchor.addprevious(element)

    enable_field_updates(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def write_manifest(output_docx: Path) -> None:
    manifest_path = APPENDIX_ROOT / "appendix_manifest.csv"
    rows = []
    for path in sorted(APPENDIX_ROOT.rglob("*")):
        if not path.is_file() or path == manifest_path or path.suffix.lower() in {".pyc", ".ndjson"}:
            continue
        relative = path.relative_to(APPENDIX_ROOT)
        appendix = relative.parts[0][0] if relative.parts and re.match(r"^[A-G] ", relative.parts[0]) else "Root"
        digest = hashlib.sha256(path.read_bytes()).hexdigest()
        rows.append([appendix, path.suffix.lower().lstrip("."), str(relative).replace("\\", "/"), path.stem, path.stat().st_size, digest])
    if output_docx.exists():
        digest = hashlib.sha256(output_docx.read_bytes()).hexdigest()
        rows.append(["Integrated manuscript", "docx", output_docx.name, output_docx.stem, output_docx.stat().st_size, digest])
    with manifest_path.open("w", encoding="utf-8-sig", newline="") as stream:
        writer = csv.writer(stream)
        writer.writerow(["appendix", "file_type", "file", "title", "bytes", "sha256"])
        writer.writerows(rows)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--manifest-only", action="store_true")
    return parser.parse_args()


if __name__ == "__main__":
    arguments = parse_args()
    if not arguments.manifest_only:
        if arguments.source is None:
            raise SystemExit("--source is required unless --manifest-only is used")
        assemble(arguments.source, arguments.output)
    write_manifest(arguments.output)
    print(arguments.output)
